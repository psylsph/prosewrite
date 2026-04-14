from __future__ import annotations

import re
from pathlib import Path

import questionary
from rich.table import Table

from ..approval import _STYLE
from ..config import ProjectConfig
from ..display import console, show_info, show_success, show_warning, word_count
from ..state import ProjectState, save_state


def _chapter_num(p: Path) -> int:
    m = re.search(r"chapter_(\d+)", p.name)
    return int(m.group(1)) if m else 9999


def _md_to_html(text: str) -> str:
    """Convert markdown chapter text to HTML for EPUB."""
    import markdown as md_lib
    return md_lib.markdown(text, extensions=["extra", "nl2br"])


# ── Format writers ────────────────────────────────────────────────────────────

def _write_markdown(
    project_dir: Path,
    project_name: str,
    manuscript_parts: list[tuple[str, str]],
    total_words: int,
) -> Path:
    """Write manuscript.md — clean markdown with chapter headings."""
    sections = [f"# {project_name}\n"]
    for heading, body in manuscript_parts:
        sections.append(f"# {heading}\n\n{body.strip()}")
    text = "\n\n---\n\n".join(sections)
    path = project_dir / "manuscript.md"
    path.write_text(text, encoding="utf-8")
    show_success(f"manuscript.md written ({total_words:,} words).")
    return path


def _write_docx(
    project_dir: Path,
    project_name: str,
    manuscript_parts: list[tuple[str, str]],
) -> Path | None:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        show_warning("python-docx not installed — skipping .docx export.")
        return None

    doc = Document()

    # Title page
    title_para = doc.add_heading(project_name, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for heading, body in manuscript_parts:
        doc.add_page_break()
        doc.add_heading(heading, level=1)
        for line in body.splitlines():
            line = line.strip()
            if not line:
                continue
            if line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            else:
                doc.add_paragraph(line)

    path = project_dir / "manuscript.docx"
    doc.save(str(path))
    show_success(f"manuscript.docx written.")
    return path


def _write_epub(
    project_dir: Path,
    project_name: str,
    manuscript_parts: list[tuple[str, str]],
    author: str = "",
) -> Path | None:
    try:
        from ebooklib import epub
    except ImportError:
        show_warning("ebooklib not installed — skipping .epub export.")
        return None

    book = epub.EpubBook()
    book.set_title(project_name)
    book.set_language("en")
    if author:
        book.add_author(author)

    epub_chapters: list = []
    toc_entries: list = []

    for i, (heading, body) in enumerate(manuscript_parts, 1):
        filename = f"chap_{i:03d}.xhtml"
        html_body = _md_to_html(body)
        content = (
            f"<html><body>"
            f"<h1>{heading}</h1>"
            f"{html_body}"
            f"</body></html>"
        )
        chap = epub.EpubHtml(title=heading, file_name=filename, lang="en")
        chap.content = content
        book.add_item(chap)
        epub_chapters.append(chap)
        toc_entries.append(epub.Link(filename, heading, f"chap{i}"))

    book.toc = tuple(toc_entries)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    # Basic stylesheet
    css = epub.EpubItem(
        uid="style",
        file_name="style/main.css",
        media_type="text/css",
        content=b"body { font-family: Georgia, serif; line-height: 1.6; margin: 2em; }"
                b"h1 { margin-top: 2em; } p { text-indent: 1.5em; margin: 0; }",
    )
    book.add_item(css)
    for chap in epub_chapters:
        chap.add_item(css)

    book.spine = ["nav"] + epub_chapters

    path = project_dir / "manuscript.epub"
    epub.write_epub(str(path), book, {})
    show_success(f"manuscript.epub written.")
    return path


# ── Stage entry point ─────────────────────────────────────────────────────────

def run(cfg: ProjectConfig, project_dir: Path, state: ProjectState) -> None:
    """Stage 6 — Export approved chapters to selected formats."""
    chapters_dir = project_dir / "chapters"
    chapter_files = (
        sorted(chapters_dir.glob("chapter_*.md"), key=_chapter_num)
        if chapters_dir.exists() else []
    )

    if not chapter_files:
        show_warning("No approved chapters found in chapters/ directory.")
        return

    # Build manuscript parts and summary table
    manuscript_parts: list[tuple[str, str]] = []
    table = Table(title="Manuscript", show_lines=True)
    table.add_column("#", style="bold", justify="right")
    table.add_column("File")
    table.add_column("Words", justify="right")
    total_words = 0

    for chapter_file in chapter_files:
        text = chapter_file.read_text(encoding="utf-8")
        wc = word_count(text)
        total_words += wc
        m = re.search(r"chapter_(\d+)", chapter_file.name)
        num = m.group(1) if m else "?"
        label = f"Chapter {num}" if num != "0" else "Prelude"
        table.add_row(num, chapter_file.name, f"{wc:,}")
        manuscript_parts.append((label, text))

    table.add_section()
    table.add_row("", "[bold]TOTAL[/bold]", f"[bold]{total_words:,}[/bold]")
    console.print(table)

    # Format selection
    print()
    formats = questionary.checkbox(
        "Select export formats",
        choices=[
            questionary.Choice("Markdown  (.md)",  value="md",   checked=True),
            questionary.Choice("EPUB      (.epub)", value="epub", checked=True),
            questionary.Choice("Word      (.docx)", value="docx", checked=False),
        ],
        style=_STYLE,
    ).ask()

    if not formats:
        show_warning("No formats selected — nothing exported.")
        return

    if "md" in formats:
        _write_markdown(project_dir, state.project_name, manuscript_parts, total_words)

    if "epub" in formats:
        author = getattr(cfg, "author", "") or ""
        _write_epub(project_dir, state.project_name, manuscript_parts, author)

    if "docx" in formats:
        _write_docx(project_dir, state.project_name, manuscript_parts)

    state.current_stage = "stage6_export"
    save_state(state, project_dir)
